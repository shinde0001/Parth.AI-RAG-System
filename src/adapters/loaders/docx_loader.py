import uuid
from pathlib import Path

import docx

from src.core.models.document import Document
from src.core.ports.document_loader_port import DocumentLoaderPort


class DocxLoaderAdapter(DocumentLoaderPort):
    """Adapter for loading and parsing DOCX documents."""
    
    def load(self, file_path: str) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        doc = docx.Document(str(path))
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        return Document(
            document_id=str(uuid.uuid4()),
            filename=path.name,
            content=content,
            metadata={"source_type": "docx"}
        )
